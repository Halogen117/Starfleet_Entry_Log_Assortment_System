import tkinter
from Stardate.sdcompute import Compute
from Stardate.sdview import View
from PIL import ImageTk,Image
import os
import pyglet
import getpass
import re
from docx import Document
import docx

class start_program:
    def __init__(self):
        print("Initializing!")
        self.one_true_window = True


    def center_window(self,window, w=300, h=200):
        # get screen width and height
        ws = window.winfo_screenwidth()
        hs = window.winfo_screenheight()
        # calculate position x, y
        x = (ws/2) - (w/2)    
        y = (hs/2) - (h/2)
        window.geometry('%dx%d+%d+%d' % (w, h, x, y))

    def home_window(self):
        self.interface = tkinter.Tk()
        self.interface.title("LCARS *Stardate*")
        self.interface.configure(bg='black')
        self.center_window(self.interface,1200,600)

        pyglet.font.add_file(os.getcwd()+"\\okuda\\Okuda.otf") 

        top_frame = tkinter.Frame(self.interface).pack()
        bottom_frame = tkinter.Frame(self.interface).pack(side = "bottom")
        left_frame = tkinter.Frame(self.interface).pack(side = "left")
        middle_frame = tkinter.Frame(self.interface).pack(fill="none", expand=True)
        right_frame = tkinter.Frame(self.interface).pack(side="left")

        btn1 = tkinter.Label(top_frame,text = "Personnel File", fg = "red", font=('Okuda',25))

        pil_img = Image.open(os.getcwd()+"\\i_tried.jpg")
        img_file = ImageTk.PhotoImage(pil_img)
        canvas = tkinter.Canvas(middle_frame,width= pil_img.width, height=pil_img.height)
        canvas.create_image(10,10,image=img_file,anchor=tkinter.NW)


        pil_img_LCARS = Image.open(os.getcwd()+"\\top_image.gif")
        img_file_LCARS  = ImageTk.PhotoImage(pil_img_LCARS)
        canvas_LCARS = tkinter.Canvas(middle_frame,width=pil_img_LCARS.width, height=pil_img_LCARS.height)
        canvas_LCARS.create_image(10,10,image=img_file_LCARS,anchor=tkinter.NW)



        label_name = tkinter.Label(bottom_frame,text = "Name of Individual : Russel Tan Jun Hong", font=('Okuda',25))
        label_species = tkinter.Label(right_frame,text = "Species : Human",  font=('Okuda',25))
        label_birth = tkinter.Label(right_frame,text = "Born : 2000",  font=('Okuda',25))
        label_rank = tkinter.Label(right_frame,text = "Rank : Recruit/ Date of Rank : 98513.2",  font=('Okuda',25))
        label_notes = tkinter.Label(right_frame,text = "Notes : Achieved the CEH Certification",  font=('Okuda',25))
        label_notes_2 = tkinter.Label(right_frame,text = "      Earned the Vulcan Medal of Surak",  font=('Okuda',25))


        btn = tkinter.Button(self.interface,  
                    text ="Click to open a new window",  
                    command = lambda: openWindowLog(self)) 
        btn.pack(pady = 10) 

        btn1.pack()
        canvas.pack(side = "left")
        canvas_LCARS.pack()
        label_name.pack()
        label_species.pack()
        label_birth.pack()
        label_rank.pack()
        label_notes.pack()
        label_notes_2.pack()
        self.interface.mainloop()

def destroy_window(windows):
    windows.destroy()

def record_log_window(self):
    print("Oppps")

def read_log_window(self):
    # Access Common file directory and get the word documents there (OS Module)
    current_user = getpass.getuser()
    store_retrieved_stardates = []
    for root, dirs, files in os.walk('C:\\Users\\%s\\Documents\\Log_Entry' %current_user):
        for filename in files:
            store_retrieved_stardates.append(re.sub('.docx','',filename))
    
    if View.return_today_stardate()['Stardate'] is not None:
        print("These are the available Stardates that can be reviewed. Please select the stardate you want to read")
        stardate_info_label = tkinter.Label(self.personal_log, text = "Stardates : ", font=('Okuda',25))
        for stardates in store_retrieved_stardates:
            stardate_label =  tkinter.Button(self.personal_log, text = stardates, font=('Okuda',25))
            stardate_label.pack()

            result = docx.Document('C:\\Users\\'+current_user+'\\Documents\\Log_Entry\\'+ stardates+".docx")
            result = [p.text for p in result.paragraphs]
            personal_log_label = tkinter.Label(self.personal_log, text = result, font=('Okuda',25))
            personal_log_label.pack()

    else:
        # Retrieve text from log
        result = docx.Document('C:\\Users\\'+current_user+'\\Documents\\Log_Entry\\'+ View.return_today_stardate()['Stardate']+".docx")
        result = [p.text for p in result.paragraphs]
        print(result)

def openWindowLog(self):
    # Toplevel object which will  
    # be treated as a new window 
    self.personal_log = tkinter.Tk()
    # sets the title of the 
    # Toplevel widget 
    self.personal_log.title("LCARS *Stardate*") 
  
    # sets the geometry of toplevel 
    
    self.center_window(self.personal_log,1200,600)
  
    # A Label widget to show in toplevel 
    tkinter.Label(self.personal_log,  text ="Personeel Log",font=('Okuda',25)).pack() 

    btn_record_log = tkinter.Button(self.personal_log,  
                text ="Record Log",  
                command = lambda: record_log_window(self)) 
    
    btn_read_log = tkinter.Button(self.personal_log,  
                text ="Read Log",  
                command = lambda: read_log_window(self)) 

    btn_back_to_self= tkinter.Button(self.personal_log,  text ="Back",  command = lambda: self.home_window()) 

    btn_read_log.pack() 
    btn_record_log.pack()
    btn_back_to_self.pack()
    destroy_window(self.interface)



def main():
    program_start = start_program()
    program_start.home_window()

main()