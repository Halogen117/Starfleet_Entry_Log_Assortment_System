# Credits to the dudes that made the python code for the stardate calculator!
# Check em out
import argparse
import speech_recognition as sr
from Stardate.sdcompute import Compute
from Stardate.sdview import View
from docx import Document
import docx
import os
import getpass
import re
from pynput import keyboard


global break_program

def on_press(key):
    global break_program
    print (key)
    if key == keyboard.Key.esc:
        print ('end pressed')
        break_program = True
        return False

def record_log():
    global break_program
    break_program = False
    converted_text = []
    
    # Establish saved directory
    current_user = getpass.getuser()
    if not os.path.exists('C:\\Users\\%s\\Documents\\Log_Entry' %current_user):
        os.chdir('C:\\Users\\%s\\Documents' %current_user)
        os.mkdir('Log_Entry')
    listen_audio = sr.Recognizer()

    print("Welcome to the Starfleet text to audio converter!")
    return_stardate_dict =  View.return_today_stardate()

    # Establish stardate
    docx_file = 'C:\\Users\\'+current_user+'\\Documents\\Log_Entry\\'+return_stardate_dict['Stardate']+'.docx'

    print("Today's Stardate is "+return_stardate_dict['Stardate']+" for reference purpose")
    if os.path.exists(docx_file):
        print("Note there is already a log on the same stardate! This will be considered a supplemental log.")

    # Record sample log (Debug Microphone)
    saved_wav_file = return_stardate_dict['Stardate']+'.wav'

    #print(sr.Microphone.list_microphone_names())
    my_mic = sr.Microphone(device_index=1) #my device index is 1, you have to put your device index
    with keyboard.Listener(on_press=on_press) as listener:
        while break_program == False:
            with my_mic as source:
                print("Beginning audio recording")
                listen_audio.adjust_for_ambient_noise(source)
                audio = listen_audio.listen(source, timeout=None)

            # Convert the audio to text. With the new lines of code coming, should be saved in an array? 
            converted_text.append(listen_audio.recognize_google(audio))
            print(converted_text)
    listener.join()
    
    # Save the file
    if not os.path.exists(docx_file):
        document = Document()
        document.add_heading('Stardate '+return_stardate_dict['Stardate'], level=1)
        for i in range(len(converted_text)):
            document.add_paragraph(converted_text[i])
    else:
        # Supplemental logs
        document = Document(docx_file)
        document.add_heading('Supplemental Log', level=1)
        for i in range(len(converted_text)):
            document.add_paragraph(converted_text[i])

    document.save(docx_file)

    # Save the file
    print("done - result written to "+saved_wav_file)


def read_log(args_stardate_specify):
    # Access Common file directory and get the word documents there (OS Module)
    current_user = getpass.getuser()
    store_retrieved_stardates = []
    for root, dirs, files in os.walk('C:\\Users\\%s\\Documents\\Log_Entry' %current_user):
        for filename in files:
            store_retrieved_stardates.append(re.sub('.docx','',filename))
    
    if args_stardate_specify is None:
        print("These are the available Stardates that can be reviewed. Please select the stardate you want to read")
        for stardates in store_retrieved_stardates:
            print(stardates)
    else:
        # Retrieve text from log
        result = docx.Document('C:\\Users\\'+current_user+'\\Documents\\Log_Entry\\'+ args_stardate_specify+".docx")
        result = [p.text for p in result.paragraphs]
        print(result)
def main():
    parser = argparse.ArgumentParser(description='Welcome to the voice to text Starfleet prototype! Record and read logs with style!')
    parser.add_argument('--read', action='store_false', dest='read_log',
                    help='Specify this option for the program to read text from a file')
    parser.add_argument('--record', action='store_true', dest='record_log',
                    help='Specify this option for the program to record your text into a file')
    parser.add_argument('--stardate', action='store', dest='stardate_specify',
                    help='Add in the stardate you want to read!', default=None, type=str)

    args = parser.parse_args()
    args_read_log = args.read_log
    args_record_log = args.record_log
    args_stardate_specify = args.stardate_specify
    
    if args_read_log is False:
        read_log(args_stardate_specify)
    elif args_record_log is True:
        record_log()
    else:
        print("Welcome to the Starfleet voice_to_text prototype!")
        print("Do you wish to check logs or record a new log?")
        print("1 - Record Log")
        print("2 - Read Log")
        input_take = input("Insert Reply Here --> ")
        if input_take == "1":
            record_log()
        else:
            read_log(args_stardate_specify)





main()